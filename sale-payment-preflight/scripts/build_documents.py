#!/usr/bin/env python3
"""Build the two lawyer-review DOCX deliverables from analysis.json."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DISCLAIMER = "律师复核前草稿：仅供案卷预整理，不构成法律意见，不得替代律师独立判断。"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RED = "9B1C1C"
GOLD = "7A5A00"
MUTED = "667085"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    set_run_font(header_p.add_run(running_label), size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    set_run_font(footer_p.add_run("律师复核前草稿  |  第 "), size=8.5, color=MUTED)
    run = footer_p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])
    set_run_font(run, size=8.5, color=MUTED)
    set_run_font(footer_p.add_run(" 页"), size=8.5, color=MUTED)

    props = doc.core_properties
    props.author = "sale-payment-preflight"
    props.last_modified_by = "sale-payment-preflight"
    props.subject = "买卖合同货款追索案卷体检"
    props.keywords = "合成案卷, 律师复核, 买卖合同, 货款追索"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_table(table, widths_dxa: list[int], header: bool = True) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"表格列宽合计必须为 {TABLE_WIDTH_DXA} DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for name, attrs in (
        ("tblW", {"w": str(TABLE_WIDTH_DXA), "type": "dxa"}),
        ("tblInd", {"w": str(TABLE_INDENT_DXA), "type": "dxa"}),
        ("tblLayout", {"type": "fixed"}),
    ):
        element = tbl_pr.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tbl_pr.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), value)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "D9DEE7")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0 and header:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[col_index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[col_index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0 and header:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=(row_index == 0 and header), color=(NAVY if row_index == 0 and header else None))


def add_title_block(doc: Document, title: str, subtitle: str, analysis: dict[str, Any]) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    set_run_font(kicker.add_run("企业买卖合同货款追索 | 诉前工作底稿"), size=9.5, bold=True, color=BLUE)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(5)
    set_run_font(title_p.add_run(title), size=25, bold=True, color=NAVY)
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle_p.add_run(subtitle), size=12.5, color=MUTED)
    case = analysis["case"]
    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("案件", case["case_name"]),
        ("委托方", case["client_name"]),
        ("相对方", case["counterparty_name"]),
        ("截止日 / 币种", f"{case['as_of_date']} / {case['currency']}"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    configure_table(table, [2700, 6660], header=False)
    doc.add_paragraph()
    add_callout(doc, DISCLAIMER)


def add_callout(doc: Document, text: str, fill: str = CALLOUT, color: str = NAVY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE)
    borders.append(left)
    p_pr.append(borders)
    set_run_font(paragraph.add_run(text), size=10.5, bold=True, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, bold_label: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_label:
        set_run_font(paragraph.add_run(bold_label), bold=True, color=NAVY)
    set_run_font(paragraph.add_run(text), size=11)


def severity_zh(value: str) -> str:
    return {"high": "高", "medium": "中", "low": "低"}.get(value, value)


def create_issue_document(analysis: dict[str, Any], output: Path) -> None:
    doc = Document()
    configure_document(doc, "买卖合同货款追索 | 矛盾与补证清单")
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.15
    add_title_block(doc, "矛盾与补证清单", "按风险等级排列的律师复核工作表", analysis)
    counts = {level: sum(1 for issue in analysis["issues"] if issue["severity"] == level) for level in ("high", "medium", "low")}
    add_body(doc, f"共发现 {len(analysis['issues'])} 个待复核事项，其中高风险 {counts['high']} 个、中风险 {counts['medium']} 个、低风险 {counts['low']} 个。", "审查摘要：")
    add_heading(doc, "问题清单", 1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["等级", "问题及说明", "来源", "补证/复核建议"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for issue in analysis["issues"]:
        cells = table.add_row().cells
        cells[0].text = severity_zh(issue["severity"])
        cells[1].text = f"{issue['id']}｜{issue['title']}\n{issue['detail']}"
        cells[2].text = ", ".join(issue["source_ids"]) or "待定位"
        cells[3].text = issue["recommendation"]
        if issue["severity"] == "high":
            set_cell_shading(cells[0], "FCE8E6")
        elif issue["severity"] == "medium":
            set_cell_shading(cells[0], "FFF4CE")
    configure_table(table, [900, 3500, 1200, 3760])
    add_heading(doc, "律师复核签记", 1)
    review = doc.add_table(rows=5, cols=2)
    rows = [
        ("第一轮复核人/日期", ""),
        ("已纠正的事实或金额", ""),
        ("仍需向客户补充的问题", ""),
        ("第二轮复核人/日期", ""),
        ("结论", "□ 通过  □ 退回修改  □ 超出本 Skill 范围"),
    ]
    for row, (label, value) in zip(review.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    configure_table(review, [2700, 6660], header=False)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def create_report_document(analysis: dict[str, Any], output: Path) -> None:
    doc = Document()
    configure_document(doc, "买卖合同货款追索 | 诉前案卷体检报告")
    add_title_block(doc, "诉前案卷体检报告", "材料、履约、金额和证据链的可追溯预审", analysis)

    add_heading(doc, "1. 范围与方法", 1)
    add_body(doc, "本报告仅覆盖境内 B2B 货物买卖中卖方追索未付货款。系统只读建立材料清单和哈希，抽取可定位事实线索，以已确认 CNY 流水复算未付本金，并按规则矩阵提示缺证和矛盾。")
    add_body(doc, "本报告不判断胜诉可能性，不给出最终时效、管辖或仲裁结论，不计算利息、违约金或诉讼费用。")

    add_heading(doc, "2. 核心摘要", 1)
    high_count = sum(1 for issue in analysis["issues"] if issue["severity"] == "high")
    blocked = analysis["audit"]["unsupported_or_blocked_files"]
    summary = doc.add_table(rows=2, cols=4)
    summary_rows = [
        ("材料数", str(len(analysis["materials"])), "可追溯事实", str(len(analysis["facts"]))),
        ("高风险事项", str(high_count), "未付本金（CNY）", analysis["ledger"]["outstanding_principal"]),
    ]
    for row, values in zip(summary.rows, summary_rows):
        for index, value in enumerate(values):
            row.cells[index].text = value
            if index % 2 == 0:
                set_cell_shading(row.cells[index], LIGHT_GRAY)
    configure_table(summary, [1700, 1700, 2600, 3360], header=False)
    if blocked:
        add_callout(doc, f"有 {blocked} 份材料未完成文本抽取，相关内容不得推断。", fill="FFF4CE", color=GOLD)

    add_heading(doc, "3. 案件材料", 1)
    material_table = doc.add_table(rows=1, cols=4)
    for cell, value in zip(material_table.rows[0].cells, ["编号", "文件", "状态", "完整性/安全备注"]):
        cell.text = value
    for material in analysis["materials"]:
        cells = material_table.add_row().cells
        cells[0].text = material["id"]
        cells[1].text = material["relative_path"]
        cells[2].text = material["status"]
        flags = []
        if material["embedded_instruction_detected"]:
            flags.append("疑似嵌入指令")
        if sum(material["pii_counts"].values()):
            flags.append(f"敏感字段 {sum(material['pii_counts'].values())} 处")
        if material["note"]:
            flags.append(material["note"])
        cells[3].text = "；".join(flags) or "已建立哈希和位置索引"
    configure_table(material_table, [800, 3800, 1400, 3360])

    add_heading(doc, "4. 履约事实时间轴", 1)
    if analysis["timeline"]:
        timeline_table = doc.add_table(rows=1, cols=4)
        for cell, value in zip(timeline_table.rows[0].cells, ["日期", "类别", "事件摘要", "来源"]):
            cell.text = value
        for event in analysis["timeline"][:30]:
            cells = timeline_table.add_row().cells
            cells[0].text = event["date"]
            cells[1].text = event["category"]
            cells[2].text = event["event"]
            cells[3].text = f"{event['material_id']} / {event['location']} / {event['fact_id']}"
        configure_table(timeline_table, [1400, 1400, 3900, 2660])
    else:
        add_callout(doc, "未抽取到可定位日期；时间轴待人工补充。", fill="FFF4CE", color=GOLD)

    add_heading(doc, "5. 货款复算", 1)
    ledger = analysis["ledger"]
    ledger_table = doc.add_table(rows=5, cols=2)
    amount_rows = [
        ("已确认应收", ledger["totals"]["应收"]),
        ("已确认付款", ledger["totals"]["付款"]),
        ("已确认退货", ledger["totals"]["退货"]),
        ("已确认折让", ledger["totals"]["折让"]),
        ("未付本金", ledger["outstanding_principal"]),
    ]
    for row, (label, value) in zip(ledger_table.rows, amount_rows):
        row.cells[0].text = label
        row.cells[1].text = f"CNY {value}"
        set_cell_shading(row.cells[0], LIGHT_GRAY)
    configure_table(ledger_table, [3600, 5760], header=False)
    for row in ledger_table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    add_body(doc, f"计入 {ledger['included_rows']} 行；排除 {ledger['excluded_rows']} 行。公式：应收 - 付款 - 退货 - 折让。", "复算口径：")

    add_heading(doc, "6. 要件—证据矩阵", 1)
    matrix = doc.add_table(rows=1, cols=4)
    for cell, value in zip(matrix.rows[0].cells, ["规则", "审查事项", "状态", "证据线索/复核提示"]):
        cell.text = value
    for item in analysis["evidence_matrix"]:
        cells = matrix.add_row().cells
        cells[0].text = f"{item['rule_id']}\n{item['category']}"
        cells[1].text = item["question"]
        cells[2].text = item["status"]
        evidence = ", ".join(item["evidence_fact_ids"]) or "无"
        cells[3].text = f"事实：{evidence}\n复核：{item['lawyer_note']}"
    configure_table(matrix, [1200, 3000, 1500, 3660])

    add_heading(doc, "7. 关键矛盾与补证", 1)
    for issue in analysis["issues"][:12]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        color = RED if issue["severity"] == "high" else GOLD if issue["severity"] == "medium" else MUTED
        set_run_font(paragraph.add_run(f"[{severity_zh(issue['severity'])}] {issue['id']} {issue['title']}："), bold=True, color=color)
        set_run_font(paragraph.add_run(f"{issue['detail']} 建议：{issue['recommendation']} 来源：{', '.join(issue['source_ids']) or '待定位'}。"), size=10.5)

    add_heading(doc, "8. 律师复核门禁", 1)
    checklist = doc.add_table(rows=7, cols=2)
    checks = [
        ("□", "确认案件属于本 Skill 范围"),
        ("□", "确认主体、签章、授权、收付款主体"),
        ("□", "复核交付、验收、质量异议和退换货"),
        ("□", "逐行复核货款流水及未付本金"),
        ("□", "独立判断管辖/仲裁和时效"),
        ("□", "核验法源版本、条文及案件时间效力"),
        ("□", "决定成果是否可发送或提交第三方"),
    ]
    for row, values in zip(checklist.rows, checks):
        row.cells[0].text, row.cells[1].text = values
    configure_table(checklist, [800, 8560], header=False)
    for row in checklist.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=40, bottom=40)

    add_heading(doc, "9. 法源登记", 1)
    sources = {
        "L01": "《中华人民共和国民法典》 https://wb.flk.npc.gov.cn/flfg/PDF/bd53dd912c1048f2aecbaa229238334b.pdf",
        "L02": "买卖合同司法解释（2020 修正） https://gongbao.court.gov.cn/Details/ba1b35d94fe5430b29f99395323878.html",
        "L03": "民事诉讼证据规定（2019 修正） https://www.court.gov.cn/zixun/xiangqing/212721.html",
        "L04": "《中华人民共和国律师法》 https://www.npc.gov.cn/npc/c2/c183/c198/201905/t20190522_27500.html",
    }
    for source_id in analysis.get("legal_sources", []):
        if source_id in sources:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.05
            set_run_font(paragraph.add_run(f"{source_id}："), size=9.5, bold=True, color=NAVY)
            set_run_font(paragraph.add_run(sources[source_id]), size=9.5)
    add_callout(doc, "法源核验日期为 2026-08-12。外部使用前须由承办律师重新确认有效性、时间效力和案件适用范围。")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--analysis", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    args = parser.parse_args()
    analysis = json.loads(args.analysis.read_text(encoding="utf-8"))
    args.output_dir.mkdir(parents=True, exist_ok=True)
    create_issue_document(analysis, args.output_dir / "05_矛盾与补证清单.docx")
    create_report_document(analysis, args.output_dir / "06_诉前案卷体检报告.docx")
    print(json.dumps({"status": "ok", "documents": 2, "output_dir": str(args.output_dir.resolve())}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
