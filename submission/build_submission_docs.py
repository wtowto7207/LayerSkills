#!/usr/bin/env python3
"""Build the competition-facing technical specification document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "submission" / "参赛作品技术说明书.docx"
PREFILL_OUTPUT = ROOT / "submission" / "报名信息预填稿_非官方模板.docx"
NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "667085"
GOLD = "7A5A00"
RED = "9B1C1C"
TABLE_WIDTH = 9360


def font(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    rfonts = run._element.get_or_add_rPr().rFonts
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    element = tc_pr.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        tc_pr.append(element)
    element.set(qn("w:fill"), fill)


def margins(cell, value=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value if side in {"top", "bottom"} else 120))
        node.set(qn("w:type"), "dxa")


def table_style(table, widths, header=True):
    if sum(widths) != TABLE_WIDTH:
        raise ValueError("表格列宽合计错误")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), "120")
    tbl_indent.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "D9DEE7")
        borders.append(node)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0 and header:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index] / 1440)
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths[col_index]))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            if row_index == 0 and header:
                shade(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1.5)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    font(run, 9.2, row_index == 0 and header, NAVY if row_index == 0 and header else None)


def setup(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(5)
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 12.5, 12, 5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    font(header.add_run("浙江省青年律师法律 AI Skill 征集｜作品技术说明"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("sale-payment-preflight  |  第 "), 8.5, color=MUTED)
    run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    font(run, 8.5, color=MUTED)
    font(footer.add_run(" 页"), 8.5, color=MUTED)
    doc.core_properties.author = "sale-payment-preflight 参赛团队"
    doc.core_properties.subject = "企业买卖合同货款追索诉前案卷体检 Skill"


def p(doc, text, label=None, size=10.5, color=None):
    paragraph = doc.add_paragraph()
    if label:
        font(paragraph.add_run(label), size, True, NAVY)
    font(paragraph.add_run(text), size, color=color)
    return paragraph


def list_item(doc, text, numbered=False, size=10.5, style_name=None):
    paragraph = doc.add_paragraph(style=style_name or ("List Number" if numbered else "List Bullet"))
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    font(paragraph.add_run(text), size)
    return paragraph


def callout(doc, text, fill="F4F6F9", color=NAVY):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "6")
    border.append(left)
    p_pr.append(border)
    font(paragraph.add_run(text), 10.5, True, color)


def add_rows(table, rows):
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = str(value)


def build_technical_spec():
    doc = Document()
    setup(doc)
    kicker = doc.add_paragraph()
    font(kicker.add_run("通用法律实务 Skill｜平台中立｜合成案卷演示"), 10, True, BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    font(title.add_run("企业买卖合同货款追索\n诉前案卷体检 Skill"), 25, True, NAVY)
    subtitle = doc.add_paragraph()
    font(subtitle.add_run("sale-payment-preflight｜参赛作品技术说明书"), 13, color=MUTED)
    callout(doc, "结论：作品已形成可运行 Skill 包、三套完整合成案卷、固定格式输出及自动验收；适合按“法律专业价值 + 合规安全”双主线申报。")
    info = doc.add_table(rows=4, cols=2)
    for row, values in zip(info.rows, [
        ("申报赛道", "通用法律实务 Skill"),
        ("适用场景", "境内 B2B 货物买卖中卖方追索未付货款"),
        ("版本 / 核验日", "v1.0 / 2026-08-12"),
        ("作品状态", "技术验收与 4 分钟演示已完成；律师规则确认、两轮盲测与第二智能体实跑待完成"),
    ]):
        row.cells[0].text, row.cells[1].text = values
        shade(row.cells[0], LIGHT_GRAY)
    table_style(info, [2300, 7060], header=False)
    doc.add_page_break()

    doc.add_heading("1. 作品定位与边界", 1)
    p(doc, "本 Skill 将合同、订单、送货、签收、发票、对账、付款和往来材料整理为可追溯证据链，并输出履约时间轴、要件—证据矩阵、货款复算和补证建议。所有结论均标注为“律师复核前草稿”。")
    scope = doc.add_table(rows=1, cols=3)
    for cell, value in zip(scope.rows[0].cells, ["维度", "纳入 v1", "明确排除"]):
        cell.text = value
    add_rows(scope, [
        ("主体/地域", "境内企业间交易，卖方追款", "跨境、消费、破产"),
        ("交易形态", "书面或无书面合同、分批交付、部分付款", "建工、担保、复杂质量反诉"),
        ("金额", "已确认 CNY 应收/付款/退货/折让", "多币种、利息、违约金、诉讼费"),
        ("文件", "文本型 PDF、DOCX、XLSX、CSV、TXT", "扫描件只登记并提示 OCR/人工处理"),
        ("法律判断", "呈现事实线索、缺证和矛盾", "胜诉率、最终时效/管辖/仲裁结论"),
    ])
    table_style(scope, [1500, 4000, 3860])

    doc.add_heading("2. 固定工作流与交付接口", 1)
    for index, text in enumerate([
        "只读建立文件清单、SHA-256 和材料编号，前后复核原件未变。",
        "离线提取主体、合同、交付、验收、开票、付款、催告等事实，保留材料编号与页码/工作表/行号。",
        "形成履约时间轴，并按已确认 CNY 流水精确复算未付本金。",
        "按 14 条律师规则矩阵检查合同关系、交付验收、结算、争议解决、时效事实和电子证据。",
        "将冲突、缺失和无法判断事项统一标记为待核实，进入律师复核门禁。",
    ], 1):
        list_item(doc, text, numbered=True)
    outputs = doc.add_table(rows=1, cols=2)
    outputs.rows[0].cells[0].text = "固定文件"
    outputs.rows[0].cells[1].text = "用途"
    add_rows(outputs, [
        ("01_案件材料清单.xlsx", "格式、哈希、抽取状态、敏感字段和安全备注"),
        ("02_履约事实时间轴.xlsx", "日期、事实摘要、材料编号和精确位置"),
        ("03_要件证据矩阵.xlsx", "14 条规则、状态、事实编号、法源与律师提示"),
        ("04_货款核对表.xlsx", "公式驱动的应收、已付、退货、折让和未付本金"),
        ("05_矛盾与补证清单.docx", "按风险排序的复核工作表与两轮签记"),
        ("06_诉前案卷体检报告.docx", "材料、时间轴、复算、矩阵、补证和复核门禁"),
        ("audit.json", "材料哈希、规则版本、警告、脱敏事件和复核状态"),
    ])
    table_style(outputs, [3300, 6060])

    doc.add_heading("3. 合规与安全设计", 1)
    security = doc.add_table(rows=1, cols=3)
    for cell, value in zip(security.rows[0].cells, ["控制点", "实现", "失败策略"]):
        cell.text = value
    add_rows(security, [
        ("授权与保密", "配置须明确 authorized=true；默认离线、无联网检索", "未授权即停止"),
        ("原件保护", "只读遍历、前后哈希/大小/时间戳比对", "发现变化即审计失败"),
        ("敏感信息", "手机号、身份证号、银行卡、邮箱输出脱敏；原值仅留哈希", "不在报告复现原值"),
        ("提示词注入", "材料中的“忽略规则/上传文件”等只作为证据内容", "隔离并列为高风险"),
        ("异常文件", "加密、损坏、无文本扫描件和超范围格式显式登记", "不静默猜测、不补写事实"),
        ("人工门禁", "金额、主体、时效、管辖、法源版本须律师逐项复核", "未复核不得外发"),
    ])
    table_style(security, [1800, 4700, 2860])

    doc.add_heading("4. 已完成验证", 1)
    results = doc.add_table(rows=1, cols=6)
    for cell, value in zip(results.rows[0].cells, ["合成案卷", "材料", "事实", "问题", "未付本金", "关键验证"]):
        cell.text = value
    add_rows(results, [
        ("完整履约案", "6", "44", "1", "100,000.00", "分批交付、部分付款、来源定位"),
        ("无书面合同案", "4", "22", "6", "100,000.00", "谨慎呈现履约/结算线索"),
        ("冲突缺证案", "5", "24", "12", "100,000.00", "金额冲突、第三方付款、注入、OCR"),
    ])
    table_style(results, [1700, 800, 800, 800, 1500, 3760])
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(0)
    separator.paragraph_format.space_after = Pt(2)
    metrics = doc.add_table(rows=1, cols=4)
    for cell, value in zip(metrics.rows[0].cells, ["指标", "目标", "当前结果", "状态"]):
        cell.text = value
    add_rows(metrics, [
        ("本金核算准确率", "100%", "3/3 = 100%", "通过"),
        ("事实来源定位覆盖率", "100%", "100%", "通过"),
        ("虚构事实", "0", "0", "通过"),
        ("预设关键缺证识别率", "≥90%", "4/4 = 100%", "通过"),
        ("Skill 结构快速校验", "通过", "Skill is valid", "通过"),
        ("自动化端到端/安全测试", "通过", "6 项覆盖 3 案及失败关闭", "通过"),
        ("开放格式目录兼容", "2 种布局", "Codex/Claude Code 目录布局 2/2", "通过"),
        ("第二智能体真实运行", "2 个环境", "本机 Codex 已完成；第二平台无登录态", "待完成"),
        ("律师两轮盲测", "2 轮", "盲测包与独立答案包已生成；0/2", "待完成"),
        ("演示视频", "3–5 分钟", "4:00 / 10 页 / 1080p", "通过"),
    ])
    table_style(metrics, [3300, 1700, 2500, 1860])
    callout(doc, "自动测试通过不等于法律专业验收。references/rules.json 仍标记 pending_prc_lawyer_approval，须由浙江执业律师完成规则确认和两轮盲测。", fill="FFF4CE", color=GOLD)

    doc.add_heading("5. 评审维度对应", 1)
    scoring = doc.add_table(rows=1, cols=4)
    for cell, value in zip(scoring.rows[0].cells, ["评审维度", "权重", "作品证据", "演示重点"]):
        cell.text = value
    add_rows(scoring, [
        ("法律专业价值", "30%", "14 条规则矩阵、可追溯事实、货款精确复算", "展示无书面合同与冲突案"),
        ("实用性与易用性", "25%", "固定 7 件输出、模板化配置、律师复核门禁", "3 分钟完成导入到报告"),
        ("可推广性", "20%", "平台中立、无专有数据库、薄适配层", "说明可迁移到主流智能体"),
        ("合规与安全", "25%", "离线、只读、脱敏、注入隔离、失败关闭、audit.json", "现场展示恶意文档不执行"),
    ])
    table_style(scoring, [2000, 1000, 4060, 2300])

    doc.add_heading("6. 已完成的 4 分钟演示", 1)
    p(doc, "已生成可编辑 PPTX 和 00:04:00 MP4。画面只使用合成案卷及本地生成成果；第 9 页明确区分目录迁移校验与第二智能体真实实跑。")
    demo = doc.add_table(rows=1, cols=3)
    for cell, value in zip(demo.rows[0].cells, ["时间", "操作", "讲解要点"]):
        cell.text = value
    add_rows(demo, [
        ("0:00–0:24", "作品定位", "可审计的案卷体检，不是胜诉预测器"),
        ("0:24–1:12", "业务价值与边界", "8 类材料、7 项输出、境内 B2B 卖方追款"),
        ("1:12–2:00", "工作流与完整履约案", "只读建档、来源定位、复算未付本金 100,000"),
        ("2:00–2:48", "无书面合同与冲突缺证案", "谨慎措辞、异常暴露、注入隔离与失败关闭"),
        ("2:48–3:36", "固定交付与验收", "六份律师文档、一份审计记录；技术通过、人工门禁待完成"),
        ("3:36–4:00", "平台迁移与收束", "目录校验不冒充第二智能体实跑；律师作最终判断"),
    ])
    table_style(demo, [1500, 3000, 4860])

    doc.add_heading("7. 部署、运行与复核", 1)
    for step in [
        "复制 sale-payment-preflight 文件夹或导入其 SKILL.md。",
        "按 assets/case-config.example.json 准备配置，并将案件材料放在独立只读目录。",
        "运行核心预检，生成 analysis.json 与 audit.json。",
        "生成四份 XLSX 和两份 DOCX 成果。",
        "执行 validate_outputs.py 与 validate_portability.py。",
        "由承办律师完成规则确认、两轮盲测与签记。",
    ]:
        list_item(doc, step, numbered=True, style_name="List Number 2")
    p(doc, "技术规范公布后，仅需在 agents/ 或新增 manifest 中补充平台字段，不改变核心规则、脚本与输出合同。", "适配原则：")

    doc.add_heading("8. 法源登记与使用限制", 1)
    for source in [
        "《中华人民共和国民法典》：https://wb.flk.npc.gov.cn/flfg/PDF/bd53dd912c1048f2aecbaa229238334b.pdf",
        "买卖合同司法解释（2020 修正）：https://gongbao.court.gov.cn/Details/ba1b35d94fe5430b29f99395323878.html",
        "民事诉讼证据规定（2019 修正）：https://www.court.gov.cn/zixun/xiangqing/212721.html",
        "《中华人民共和国律师法》：https://www.npc.gov.cn/npc/c2/c183/c198/201905/t20190522_27500.html",
    ]:
        list_item(doc, source, size=9.5)
    callout(doc, "法源核验日为 2026-08-12。任何真实案件使用前，承办律师须重新确认规范有效性、时间效力、地域与案件适用范围。")

    doc.add_page_break()
    doc.add_heading("9. 投稿前待完成事项", 1)
    pending = doc.add_table(rows=1, cols=3)
    for cell, value in zip(pending.rows[0].cells, ["事项", "责任人", "完成标准"]):
        cell.text = value
    add_rows(pending, [
        ("规则矩阵专业确认", "浙江执业律师", "rules.json 改为 approved，并留姓名/日期/版本"),
        ("第一轮盲测", "未参与开发的律师", "记录漏项、错项、金额和来源定位"),
        ("修订与第二轮验收", "开发者 + 复核律师", "所有阻塞项关闭并签字"),
        ("第二智能体真实实跑", "技术负责人", "登录第二平台，以同一完整案完成核心流程并保存截图/日志"),
        ("官方报名表 Word/PDF", "申报团队", "补全队名、人员、律所、联系方式并签名"),
        ("投稿邮件", "申报团队", "附 4 分钟演示视频或链接；8 月 20 日前内部发出"),
    ])
    table_style(pending, [3000, 2100, 4260])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


def setup_prefill(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    font(header.add_run("报名信息预填稿｜仅供转录至官方报名表"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("非官方模板  |  第 "), 8.5, color=MUTED)
    page_run = footer.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([begin, instruction, end])
    font(page_run, 8.5, color=MUTED)
    font(footer.add_run(" 页"), 8.5, color=MUTED)
    doc.core_properties.author = "sale-payment-preflight 参赛团队"
    doc.core_properties.subject = "浙江省青年律师法律 AI Skill 征集报名信息预填稿（非官方）"


def field_table(doc, rows, label_width=2600):
    table = doc.add_table(rows=0, cols=2)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        shade(cells[0], LIGHT_GRAY)
    table_style(table, [label_width, TABLE_WIDTH - label_width], header=False)
    return table


def build_prefill():
    doc = Document()
    setup_prefill(doc)
    kicker = doc.add_paragraph()
    font(kicker.add_run("浙江省青年律师法律 AI 技能（Skill）征集活动"), 10, True, BLUE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    font(title.add_run("报名信息预填稿"), 25, True, NAVY)
    subtitle = doc.add_paragraph()
    font(subtitle.add_run("sale-payment-preflight｜非官方模板｜2026-08-12"), 12.5, color=MUTED)
    callout(
        doc,
        "重要：本文件依据活动通知整理，不是浙江律协发布的《报名表.doc》。取得官方附件后，须逐项转录并以官方字段、声明和签名页为准；不得直接将本文件作为正式报名表投稿。",
        fill="FCE8E6",
        color=RED,
    )

    doc.add_heading("1. 参赛团队信息", 1)
    field_table(doc, [
        ("参赛队名称", "【待申报团队填写】"),
        ("牵头单位/律所", "【待填写；须符合征集对象要求】"),
        ("队长", "【姓名、执业/实习身份、所在单位】"),
        ("联系人", "【姓名】"),
        ("联系电话", "【手机号】"),
        ("电子邮箱", "【邮箱】"),
        ("联合申报单位", "【如无填“无”】"),
    ])

    doc.add_heading("2. 参赛成员与分工", 1)
    members = doc.add_table(rows=1, cols=5)
    for cell, value in zip(members.rows[0].cells, ["序号", "姓名", "身份/执业证号", "单位", "主要分工"]):
        cell.text = value
    add_rows(members, [
        ("1", "【待填写】", "【待填写】", "【待填写】", "法律规则确认、盲测与签核"),
        ("2", "【待填写】", "【待填写】", "【待填写】", "Skill 架构、脚本与兼容测试"),
        ("3", "【选填】", "【选填】", "【选填】", "材料模板、演示与投稿协调"),
    ])
    table_style(members, [700, 1500, 2100, 2200, 2860])

    doc.add_heading("3. 作品基本信息", 1)
    field_table(doc, [
        ("作品名称", "企业买卖合同货款追索诉前案卷体检 Skill"),
        ("内部名称", "sale-payment-preflight"),
        ("申报方向/赛道", "通用法律实务 Skill"),
        ("适用业务", "境内 B2B 货物买卖中卖方追索未付货款"),
        ("作品形态", "平台中立 Skill 包、确定性脚本、规则矩阵、合成案卷、固定输出模板"),
        ("版本", "v1.0（法源核验日：2026-08-12）"),
    ])

    doc.add_heading("4. 作品摘要（建议转录稿）", 1)
    p(doc, "本作品面向境内企业间货物买卖中卖方追索未付货款的诉前案卷整理。Skill 在离线、只读前提下，对合同、订单、送货签收、发票、对账、付款和往来材料建立哈希清单，提取可定位事实，形成履约时间轴、要件证据矩阵和货款核对表，并输出矛盾补证清单与诉前体检报告。对缺失、冲突、扫描件、加密文件和文档内提示词指令采取失败关闭或待核实处理，所有成果均标注为律师复核前草稿，不输出胜诉率或替代律师作终局判断。")

    doc.add_heading("5. 创新性、实用性与可推广性（建议转录稿）", 1)
    field_table(doc, [
        ("法律专业价值", "将买卖合同履约材料串为可溯源证据链，以 14 条律师规则矩阵检查合同关系、交付验收、结算、争议解决、时效事实和电子证据。"),
        ("实用与易用", "固定生成四份工作簿、两份 Word 草稿和 audit.json；支持文本型 PDF、DOCX、XLSX、CSV、TXT。"),
        ("可推广性", "核心不依赖专有法律数据库或特定模型；SKILL.md、scripts、references、assets 采用开放目录结构，平台差异由薄适配层处理。"),
        ("合规安全", "演示仅用合成案卷；默认离线、原件只读、输出脱敏、提示词注入隔离、异常文件失败关闭，并保留机器可读审计记录。"),
    ], label_width=1900)

    doc.add_heading("6. 已完成验证与仍需签核", 1)
    checks = doc.add_table(rows=1, cols=3)
    for cell, value in zip(checks.rows[0].cells, ["事项", "当前状态", "证据/完成标准"]):
        cell.text = value
    add_rows(checks, [
        ("三套合成案卷", "已完成", "完整履约、无书面合同、冲突缺证"),
        ("本金与来源验收", "已通过", "本金 3/3 正确；来源覆盖 100%；虚构事实 0"),
        ("安全失败关闭", "已通过", "未授权、买方、非 CNY 配置、输出越界、加密/不支持文件、注入与 OCR 测试"),
        ("开放格式兼容", "已通过", "两种目录布局 2/2；不等同于第二智能体真实运行"),
        ("盲测评审材料", "已完成", "评审包与独立答案包分离，避免答案泄露"),
        ("演示视频", "已完成", "00:04:00；10 页；仅使用合成材料和本地生成成果"),
        ("律师规则批准", "待完成", "浙江执业律师确认 rules.json、法源及适用边界"),
        ("律师两轮盲测", "待完成", "按《律师盲测与验收记录》填写并签记"),
        ("第二智能体实跑", "待完成", "保存同一完整案核心流程的界面截图或日志"),
    ])
    table_style(checks, [2500, 1600, 5260])

    doc.add_heading("7. 声明与签名信息准备", 1)
    p(doc, "活动通知要求申报者就材料真实性及虚假申报责任作出书面声明或承诺。以下仅为准备提示，正式文本必须原样采用官方报名表：")
    field_table(doc, [
        ("真实性承诺", "【待官方表述；全体成员核对后填写】"),
        ("知识产权/授权", "【确认 Skill、合成材料及引用内容的权属和授权】"),
        ("保密与数据合规", "【确认未提交真实客户材料或未授权信息】"),
        ("成员签名", "【取得官方表后由所有要求签名人员手写/合规电子签署】"),
        ("签署日期", "【年/月/日】"),
    ])
    callout(doc, "投稿时须同时提交：官方报名表 Word（无需签名）和全体要求人员签名后的 PDF。邮件主题格式及附件要求以主办方最新技术文档为准。", fill="FFF4CE", color=GOLD)

    PREFILL_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(PREFILL_OUTPUT)
    print(PREFILL_OUTPUT)


def build():
    build_technical_spec()
    build_prefill()


if __name__ == "__main__":
    build()
